import os
import docx
from docx.shared import Pt

def append_source_code_to_word(project_dir, docx_path):
    doc = docx.Document(docx_path)
    
    # Add page break
    doc.add_page_break()
    
    # Add heading
    doc.add_heading('PHỤ LỤC: MÃ NGUỒN CHƯƠNG TRÌNH', level=1)
    doc.add_paragraph('Dưới đây là toàn bộ mã nguồn của ứng dụng Web (ASP.NET Core) và các Script cấu hình SQL cho hệ thống TechStore.')
    
    extensions_to_include = ('.cs', '.cshtml', '.sql')
    
    for root, dirs, files in os.walk(project_dir):
        # Bỏ qua thư mục bin, obj, wwwroot
        if any(skip in root for skip in ['\\bin', '\\obj', '\\wwwroot', '\\.git', '\\docs']):
            continue
            
        for file in files:
            if file.endswith(extensions_to_include):
                file_path = os.path.join(root, file)
                rel_path = os.path.relpath(file_path, project_dir)
                
                # Heading for file
                doc.add_heading(f'File: {rel_path}', level=2)
                
                # Code block
                try:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        code_content = f.read()
                        
                    p = doc.add_paragraph(code_content)
                    p.style = 'No Spacing'
                    for run in p.runs:
                        run.font.name = 'Consolas'
                        run.font.size = Pt(10)
                except Exception as e:
                    doc.add_paragraph(f"Không thể đọc file: {e}")
                    
    doc.save(docx_path)
    print("Đã chèn toàn bộ source code vào file Word!")

if __name__ == "__main__":
    project_dir = r'c:\Users\TIEN NGUYEN\Desktop\Project'
    docx_file = r'c:\Users\TIEN NGUYEN\Desktop\Project\BaoCao_TechStore_Final.docx'
    append_source_code_to_word(project_dir, docx_file)
