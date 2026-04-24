import docx
from docx.shared import Pt
import sys

def convert_md_to_docx(md_path, docx_path):
    # Dùng file gốc làm template để giữ style (nếu có thể), nhưng an toàn nhất là tạo mới
    doc = docx.Document()
    
    # Thiết lập font mặc định
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    in_code_block = False
    in_table = False
    table_data = []

    for line in lines:
        stripped_line = line.strip()
        
        if stripped_line.startswith('```'):
            in_code_block = not in_code_block
            continue
            
        if in_code_block:
            p = doc.add_paragraph(stripped_line)
            p.style = 'No Spacing'
            for run in p.runs:
                run.font.name = 'Consolas'
                run.font.size = Pt(11)
            continue
            
        if stripped_line.startswith('|'):
            in_table = True
            row_data = [cell.strip() for cell in stripped_line.split('|')[1:-1]]
            
            # Skip divider
            if len(row_data) > 0 and all(c == '-' for c in set(row_data[0].replace(' ', ''))):
                continue
            table_data.append(row_data)
            continue
        elif in_table:
            # Render table
            if table_data and len(table_data) > 0:
                cols = len(table_data[0])
                table = doc.add_table(rows=len(table_data), cols=cols)
                table.style = 'Table Grid'
                for r_idx, row in enumerate(table_data):
                    for c_idx, cell in enumerate(row):
                        if c_idx < cols:
                            table.cell(r_idx, c_idx).text = cell
                table_data = []
            in_table = False
            
        if not stripped_line:
            doc.add_paragraph()
            continue
            
        if stripped_line.startswith('# '):
            h = doc.add_heading(stripped_line[2:], level=1)
        elif stripped_line.startswith('## '):
            h = doc.add_heading(stripped_line[3:], level=2)
        elif stripped_line.startswith('### '):
            h = doc.add_heading(stripped_line[4:], level=3)
        elif stripped_line.startswith('- '):
            p = doc.add_paragraph(stripped_line[2:], style='List Bullet')
        else:
            p = doc.add_paragraph(stripped_line)
            
    # Xử lý table ở dòng cuối cùng (nếu có)
    if in_table and table_data and len(table_data) > 0:
        cols = len(table_data[0])
        table = doc.add_table(rows=len(table_data), cols=cols)
        table.style = 'Table Grid'
        for r_idx, row in enumerate(table_data):
            for c_idx, cell in enumerate(row):
                if c_idx < cols:
                    table.cell(r_idx, c_idx).text = cell

    doc.save(docx_path)
    print("Tạo file Word thành công:", docx_path)

if __name__ == "__main__":
    md_file = r'c:\Users\TIEN NGUYEN\Desktop\Project\docs\BaoCao_TechStore_Draft.md'
    docx_file = r'c:\Users\TIEN NGUYEN\Desktop\Project\BaoCao_TechStore_Final.docx'
    convert_md_to_docx(md_file, docx_file)
